from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import google.generativeai as genai
from io import BytesIO
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi.responses import StreamingResponse

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

STRUKTURA_PROTOKOLU = [
    "Powód konsultacji:", "Aktualne samopoczucie:", "Aktywność:", "Apetyt:", "Pragnienie:",
    "Dotychczasowe żywienie:", "Smaczki i przysmaki:", "Ulubione smaki:",
    "### Kategorycznie tak:", "### Kategorycznie nie:", "### Kluczowa uwaga dot. przechowywania:",
    "Kał / Biegunka / Wymioty:", "Mocz:", "Odrobaczanie:", "Aktualne badania:", "Aktualne leki:",
    "Komentarz do wywiadu:", "Główne założenia diety:", "Co się zmieni na diecie BARF/BACF:",
    "Plan dietetyczny:", "Tranzycja i przechowywanie:", "Kaloryczność:", "Piciu:",
    "### Jakiej wody używać?", "Suplementy dodatkowe:", "Wiązanie fosforu:", "Smaczki:",
    "Inne smaczki:", "Karmy komercyjne:", "Tyndalizacja:", "Wprowadzanie suplementów:",
    "Badania kontrolne:", "Załączniki:"
]

TEKST_TYNDALIZACJA_STALY = (
    "Jeżeli robią Państwo dietę na dłużej niż 5-6 dni (mowa o diecie surowej) i chcą Państwo "
    "ją bezpiecznie przechowywać w słoiczkach w lodówce (bez zamrażania) LUB przygotowują Państwo "
    "dietę gotowaną (BACF) na zapas, konieczne jest przeprowadzenie procesu tyndalizacji (potrójnej pasteryzacji).\n\n"
    "Proces ten skutecznie eliminuje formy przetrwalnikowe bakterii (m.in. Clostridium botulinum - jadu kiełbasianego), "
    "które mogłyby namnażać się w warunkach beztlenowych zamkniętego słoika.\n\n"
    "Pełną instrukcję krok po kroku, jak prawidłowo i bezpiecznie przeprowadzić ten proces w domowych warunkach, "
    "znajdą Państwo w naszym artykule na blogu: https://meatpoint.io/pl/barf-wiedza/tyndalizacja-czyli-jak-przechowywac-posilki-jesli-nie-chcemy-ich-mrozic\n\n"
    "Dodatkowo przygotowaliśmy dla Państwa praktyczny poradnik w formie wideo na platformie YouTube, "
    "gdzie pokazujemy cały proces krok po kroku: https://www.youtube.com/watch?v=tyfT3kmq3ME"
)

TEKST_INNE_SMACZKI_STALY = (
    "Wprowadzając do codziennej rutyny jakiekolwiek inne smaczki komercyjne, należy bezwzględnie "
    "pamiętać o kontrolowaniu ich kaloryczności, aby nie zaburzyć bilansu nowej diety pacjenta.\n\n"
    "Szczegółowy poradnik oraz instrukcję, jak samodzielnie wyliczyć kaloryczność dowolnego produktu komercyjnego "
    "na podstawie danych z etykiety, znają Państwo w naszym artykule: "
    "https://meatpoint.io/pl/barf-wiedza/smaczki-i-dodatkowe-kalorie-obliczanie-kalorycznosci-komercyjnych-produktow"
)

TEKST_WPROWADZANIE_SUPLEMENTOW_STALY = (
    "Proszę zacząć od:\n• Wody\n• Mięsa\n• Podrobów\n• tłuszczu\n• Tauryny\n"
    "Proszę przygotować dietę tylko z ich zawartością i na razie pominąć pozostałe suplementy.\n\n"
    "Jak Kicia będzie się dobrze czuła, na następny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Witaminy E\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Witaminy E\n• Witamin B\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Witaminy E\n• Witamin B\n• Jodu\n• Dodatkowo: \n\n"
    "Jak wszystko będzie w porządku za kolejny tydzień proszę przygotować dietę z zawartością:\n"
    "• Wody\n• Mięsa\n• Podrobów\n• Tłuszczu / żółtka\n• Tauryny\n• Wapnia/soli\n• Kwasów omega\n• Witaminy E\n• Witamin B\n• Jodu\n• Dodatkowo: \n\n"
    "To będzie już kompletna dieta."
)

class DaneWizyty(BaseModel):
    api_key: str
    model: str
    transcript: str

def add_hyperlink(p, url, text):
    part = p.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement('w:hyperlink'); hl.set(qn('r:id'), r_id)
    nr = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    nr.append(rPr); tn = OxmlElement('w:t'); tn.text = text; nr.append(tn); hl.append(nr); p._p.append(hl)
    return hl

def parsuj_i_formatuj_tekst(p, tekst):
    parts = tekst.split('[BRAK INFORMACJI]')
    for i, part in enumerate(parts):
        if part:
            sub_segs = part.split('**')
            for idx, sub_seg in enumerate(sub_segs):
                if not sub_seg: continue
                czy_pogrubiony = (idx % 2 == 1)
                url_segs = re.split(r'(https?://[^\s]+)', sub_seg)
                for u_idx, u_seg in enumerate(url_segs):
                    if u_idx % 2 == 1:
                        add_hyperlink(p, u_seg, u_seg)
                    else:
                        run = p.add_run(u_seg)
                        if czy_pogrubiony: run.bold = True
        if i < len(parts) - 1:
            ra = p.add_run('[BRAK INFORMACJI]')
            ra.bold = True
            ra.font.color.rgb = RGBColor(220, 38, 38)

def konwertuj_do_docx(tekst_md):
    doc = Document()
    for s in doc.sections: s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Inches(1.3), Inches(0.8), Inches(0.8), Inches(0.8)
    style = doc.styles['Normal']; font = style.font; font.name, font.size, style.paragraph_format.line_spacing, style.paragraph_format.space_after = 'Arial', Pt(10.5), 1.25, Pt(4)
    
    sec = doc.sections[0]
    p_h = sec.first_page_header.paragraphs[0]
    p_h.add_run("Anna Michalska\n").bold = True
    p_h.add_run("Dietetyka Psów i Kotów\n").font.color.rgb = RGBColor(100, 116, 139)

    for line in tekst_md.split('\n'):
        l_s = line.strip()
        if not l_s: continue
        if l_s.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(l_s.replace('## ', '').replace('**', ''))
            r.bold = True; r.font.size, r.font.color.rgb = Pt(12), RGBColor(194, 65, 12)
        elif l_s.startswith('- ') or l_s.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            parsuj_i_formatuj_tekst(p, l_s.lstrip('-* '))
        else:
            p = doc.add_paragraph()
            parsuj_i_formatuj_tekst(p, l_s)
            
    b = BytesIO(); doc.save(b); return b.getvalue()

@app.post("/generuj")
async def generuj_opis(dane: DaneWizyty):
    try:
        genai.configure(api_key=dane.api_key)
        m = genai.GenerativeModel(
            model_name=dane.model,
            system_instruction="Jesteś doświadczonym, pedantycznym asystentem klinicznym dla dietetyk Anny Michalskiej. Pisz WYŁĄCZNIE prawdę na podstawie transkrypcji. Brak danych oznaczaj jako [BRAK INFORMACJI]."
        )
        
        instrukcja_szablonu = ""
        for n in STRUKTURA_PROTOKOLU:
            if n == "Tyndalizacja:": instrukcja_szablonu += f"## {n}\n{TEKST_TYNDALIZACJA_STALY}\n\n"
            elif n == "Inne smaczki:": instrukcja_szablonu += f"## {n}\n{TEKST_INNE_SMACZKI_STALY}\n\n"
            elif n == "Wprowadzanie suplementów:": instrukcja_szablonu += f"## {n}\n{TEKST_WPROWADZANIE_SUPLEMENTOW_STALY}\n\n"
            else: instrukcja_szablonu += f"## {n}\n- Uzupełnij precyzyjnie.\n"
            
        prompt = f"Przeanalizuj transkrypcję i uzupełnij szablon w tej kolejności:\nData wizyty: DD.MM.YYYY\nMetryczka pacjenta\n{instrukcja_szablonu}\n\nTranskrypcja:\n{dane.transcript}"
        
        res = m.generate_content(prompt)
        
        # 🚨 ZMIANA: Budujemy binarny plik Word w pamięci serwera
        plik_word = konwertuj_do_docx(res.text)
        
        # Odsyłamy czysty strumień pliku docx wprost do przeglądarki Vercela
        return StreamingResponse(
            BytesIO(plik_word),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Opis_Wizyty_MeatPoint.docx"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
